"""
General document processor implementation for the Limnos ingest pipeline.

This module provides a versatile implementation of the DocumentProcessor interface
that extracts metadata from various document types including text files, PDFs, 
and markdown files, using the new Pydantic models.
"""

import os
import logging
import json
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
import uuid
import re
from datetime import datetime
import mimetypes

# Document format handlers
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from limnos.ingest.processors.processor_interface import DocumentProcessor
from limnos.ingest.models.document import Document
from limnos.ingest.models.metadata import UniversalMetadata, DocumentType, Section
from limnos.pipeline.interfaces import Configurable


class GeneralDocumentProcessor(DocumentProcessor):
    """General document processor implementation that extracts metadata from various file types."""
    
    def __init__(self):
        """Initialize the general document processor."""
        self._logger = logging.getLogger(__name__)
        self._config = {}
        self._initialized = False
        self._supported_extensions = {
            '.txt': self._process_text_file,
            '.md': self._process_markdown_file,
            '.pdf': self._process_pdf_file if PYPDF_AVAILABLE else None,
            '.docx': self._process_docx_file if DOCX_AVAILABLE else None,
            '.json': self._process_json_file,
            '.html': self._process_html_file,
            '.htm': self._process_html_file,
        }
    
    def initialize(self, config: Dict[str, Any]) -> None:
        """Initialize the processor with configuration.
        
        Args:
            config: Configuration dictionary
        """
        self._config = config or {}
        
        # Set up supported extensions based on config
        if 'supported_extensions' in config:
            custom_extensions = config['supported_extensions']
            # Only keep extensions that have a corresponding processor method
            for ext in custom_extensions:
                if ext.lower() not in self._supported_extensions:
                    self._logger.warning(f"No processor available for extension {ext}")
        
        # Set up default document type mapping
        self._doc_type_mapping = config.get('doc_type_mapping', {
            '.pdf': DocumentType.PDF,
            '.txt': DocumentType.TEXT,
            '.md': DocumentType.MARKDOWN,
            '.docx': DocumentType.WORD,
            '.json': DocumentType.JSON,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
        })
        
        self._initialized = True
    
    @classmethod
    def get_plugin_name(cls) -> str:
        """Return the name of this plugin."""
        return "general"
        
    def get_default_config(self) -> Dict[str, Any]:
        """Return the default configuration for this component.
        
        Returns:
            Default configuration dictionary
        """
        return {
            'supported_extensions': ['.txt', '.md', '.pdf', '.docx', '.json', '.html', '.htm'],
            'extract_sections': True,
            'extract_metadata': True,
            'max_content_length': 10000000  # 10MB limit for content extraction
        }
    
    def process_file(self, file_path: Path) -> Document:
        """Process a file and extract content and metadata.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Processed Document with content and metadata
        """
        if not self._initialized:
            self.initialize({})
        
        try:
            extension = file_path.suffix.lower()
            
            # Check if extension is supported
            if extension in self._supported_extensions and self._supported_extensions[extension]:
                # Use the appropriate processor method
                document = self._supported_extensions[extension](file_path)
            else:
                # Fall back to binary file processing
                document = self._process_binary_file(file_path)
            
            return document
            
        except Exception as e:
            self._logger.error(f"Error processing file {file_path}: {e}")
            
            # Create a minimal document if processing fails
            doc_id = str(uuid.uuid4())
            
            # Determine document type
            extension = file_path.suffix.lower()
            doc_type = self._doc_type_mapping.get(extension, DocumentType.OTHER)
            
            # Create basic metadata
            metadata = UniversalMetadata(
                doc_id=doc_id,
                doc_type=doc_type,
                title=file_path.stem,
                file_path=str(file_path),
                storage_path=None,
                content_length=0,
                filename=file_path.name,
                extension=extension,
                size_bytes=file_path.stat().st_size,
                last_modified=datetime.fromtimestamp(file_path.stat().st_mtime),
                source_path=str(file_path),
                language='en'
            )
            
            # Create document with minimal content
            document = Document(
                doc_id=doc_id,
                content="",
                metadata=metadata,
                file_path=file_path
            )
            
            return document
    
    def _process_text_file(self, file_path: Path) -> Document:
        """Process a text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Processed Document
        """
        doc_id = str(uuid.uuid4())
        
        # Read content
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        
        # Extract basic metadata
        metadata = self._extract_basic_metadata(file_path, content, DocumentType.TEXT)
        
        # Extract text-specific metadata
        metadata.line_count = len(content.splitlines())
        metadata.word_count = len(content.split())
        
        # Try to extract sections
        if self._config.get('extract_sections', True):
            sections = self._extract_sections_from_text(content)
            if sections:
                metadata.sections = sections
        
        # Create document
        document = Document(
            doc_id=doc_id,
            content=content,
            metadata=metadata,
            file_path=file_path
        )
        
        return document
    
    def _process_markdown_file(self, file_path: Path) -> Document:
        """Process a markdown file.
        
        Args:
            file_path: Path to the markdown file
            
        Returns:
            Processed Document
        """
        doc_id = str(uuid.uuid4())
        
        # Read content
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        
        # Extract basic metadata
        metadata = self._extract_basic_metadata(file_path, content, DocumentType.MARKDOWN)
        
        # Extract markdown-specific metadata
        metadata.line_count = len(content.splitlines())
        metadata.word_count = len(content.split())
        
        # Extract title from markdown
        title = self._extract_title_from_markdown(content)
        if title:
            metadata.title = title
        
        # Extract sections from markdown
        if self._config.get('extract_sections', True):
            sections = self._extract_sections_from_markdown(content)
            if sections:
                metadata.sections = sections
        
        # Create document
        document = Document(
            doc_id=doc_id,
            content=content,
            metadata=metadata,
            file_path=file_path
        )
        
        return document
    
    def _process_pdf_file(self, file_path: Path) -> Document:
        """Process a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Processed Document
        """
        if not PYPDF_AVAILABLE:
            self._logger.warning("pypdf not available. Cannot process PDF files properly.")
            return self._process_binary_file(file_path)
        
        doc_id = str(uuid.uuid4())
        content = ""
        
        try:
            # Extract text content
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                # Extract text from all pages
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n\n"
                
                # Extract metadata from PDF info
                metadata = self._extract_basic_metadata(file_path, content, DocumentType.PDF)
                
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    
                    # Try to get title
                    if hasattr(info, 'title') and info.title:
                        metadata.title = info.title
                    
                    # Try to get authors
                    if hasattr(info, 'author') and info.author:
                        metadata.authors = [a.strip() for a in info.author.split(',')]
                    
                    # Try to get keywords
                    if hasattr(info, 'keywords') and info.keywords:
                        metadata.keywords = [k.strip() for k in info.keywords.split(',')]
                
                # Try to extract sections
                if self._config.get('extract_sections', True):
                    sections = self._extract_sections_from_text(content)
                    if sections:
                        metadata.sections = sections
        
        except Exception as e:
            self._logger.error(f"Error processing PDF {file_path}: {e}")
            # Create basic metadata
            metadata = self._extract_basic_metadata(file_path, content, DocumentType.PDF)
        
        # Create document
        document = Document(
            doc_id=doc_id,
            content=content,
            metadata=metadata,
            file_path=file_path
        )
        
        return document
    
    def _process_docx_file(self, file_path: Path) -> Document:
        """Process a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Processed Document
        """
        if not DOCX_AVAILABLE:
            self._logger.warning("python-docx not available. Cannot process DOCX files properly.")
            return self._process_binary_file(file_path)
        
        doc_id = str(uuid.uuid4())
        content = ""
        
        try:
            # Extract text content
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            content = "\n\n".join(paragraphs)
            
            # Extract basic metadata
            metadata = self._extract_basic_metadata(file_path, content, DocumentType.WORD)
            
            # Try to get document properties
            core_properties = doc.core_properties
            
            if hasattr(core_properties, 'title') and core_properties.title:
                metadata.title = core_properties.title
            
            if hasattr(core_properties, 'author') and core_properties.author:
                metadata.authors = [core_properties.author]
            
            if hasattr(core_properties, 'keywords') and core_properties.keywords:
                metadata.keywords = [k.strip() for k in core_properties.keywords.split(',')]
            
            # Extract sections from headings
            if self._config.get('extract_sections', True):
                sections = []
                current_section = None
                current_content = []
                
                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        # If we have a current section, save it
                        if current_section and current_content:
                            sections.append(Section(
                                heading=current_section,
                                content="\n\n".join(current_content),
                                level=int(current_section.style.name[-1]) if current_section.style.name[-1].isdigit() else 1
                            ))
                        
                        # Start new section
                        current_section = para.text
                        current_content = []
                    else:
                        if para.text.strip():
                            current_content.append(para.text)
                
                # Add the last section
                if current_section and current_content:
                    sections.append(Section(
                        heading=current_section,
                        content="\n\n".join(current_content),
                        level=int(current_section.style.name[-1]) if current_section.style.name[-1].isdigit() else 1
                    ))
                
                if sections:
                    metadata.sections = sections
        
        except Exception as e:
            self._logger.error(f"Error processing DOCX {file_path}: {e}")
            # Create basic metadata
            metadata = self._extract_basic_metadata(file_path, content, DocumentType.WORD)
        
        # Create document
        document = Document(
            doc_id=doc_id,
            content=content,
            metadata=metadata,
            file_path=file_path
        )
        
        return document
    
    def _process_json_file(self, file_path: Path) -> Document:
        """Process a JSON file.
        
        Args:
            file_path: Path to the JSON file
            
        Returns:
            Processed Document
        """
        doc_id = str(uuid.uuid4())
        
        try:
            # Read JSON content
            with open(file_path, 'r', encoding='utf-8') as f:
                json_content = json.load(f)
            
            # Convert to string for storage
            content = json.dumps(json_content, indent=2)
            
            # Extract basic metadata
            metadata = self._extract_basic_metadata(file_path, content, DocumentType.JSON)
            
            # Try to extract structure information
            if isinstance(json_content, dict):
                metadata.json_keys = list(json_content.keys())
            elif isinstance(json_content, list):
                metadata.json_length = len(json_content)
                if json_content and isinstance(json_content[0], dict):
                    metadata.json_keys = list(json_content[0].keys())
        
        except Exception as e:
            self._logger.error(f"Error processing JSON {file_path}: {e}")
            # Fall back to text processing
            return self._process_text_file(file_path)
        
        # Create document
        document = Document(
            doc_id=doc_id,
            content=content,
            metadata=metadata,
            file_path=file_path
        )
        
        return document
    
    def _process_html_file(self, file_path: Path) -> Document:
        """Process an HTML file.
        
        Args:
            file_path: Path to the HTML file
            
        Returns:
            Processed Document
        """
        doc_id = str(uuid.uuid4())
        
        try:
            # Read HTML content
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            
            # Extract basic metadata
            metadata = self._extract_basic_metadata(file_path, content, DocumentType.HTML)
            
            # Try to extract title
            title_match = re.search(r'<title>(.*?)</title>', content, re.IGNORECASE | re.DOTALL)
            if title_match:
                metadata.title = title_match.group(1).strip()
            
            # Extract metadata from meta tags
            description_match = re.search(r'<meta\s+name=["\'](description|Description)["\'][^>]*content=["\']([^"\']*)["\']', 
                                      content, re.IGNORECASE)
            if description_match:
                metadata.description = description_match.group(2).strip()
            
            # Extract plain text (crude approach)
            plain_text = re.sub(r'<script.*?>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
            plain_text = re.sub(r'<style.*?>.*?</style>', '', plain_text, flags=re.DOTALL | re.IGNORECASE)
            plain_text = re.sub(r'<[^>]*>', ' ', plain_text)
            plain_text = re.sub(r'\s+', ' ', plain_text).strip()
            
            # Set plain text as an attribute
            metadata.plain_text = plain_text
            
            # Try to extract sections from headings
            if self._config.get('extract_sections', True):
                sections = []
                # Find all heading tags
                heading_pattern = r'<h([1-6])[^>]*>(.*?)</h\1>'
                headings = re.finditer(heading_pattern, content, re.IGNORECASE | re.DOTALL)
                
                for match in headings:
                    level = int(match.group(1))
                    heading_text = re.sub(r'<[^>]*>', '', match.group(2)).strip()
                    
                    # Simple approach - assume content is everything until next heading of same or higher level
                    heading_end = match.end()
                    next_heading = re.search(fr'<h([1-{level}])[^>]*>', content[heading_end:], re.IGNORECASE)
                    
                    if next_heading:
                        section_content = content[heading_end:heading_end + next_heading.start()]
                    else:
                        section_content = content[heading_end:]
                    
                    # Clean up section content
                    section_content = re.sub(r'<[^>]*>', ' ', section_content)
                    section_content = re.sub(r'\s+', ' ', section_content).strip()
                    
                    sections.append(Section(
                        heading=heading_text,
                        content=section_content,
                        level=level
                    ))
                
                if sections:
                    metadata.sections = sections
                    
        except Exception as e:
            self._logger.error(f"Error processing HTML {file_path}: {e}")
            # Fall back to text processing
            return self._process_text_file(file_path)
        
        # Create document
        document = Document(
            doc_id=doc_id,
            content=content,
            metadata=metadata,
            file_path=file_path
        )
        
        return document
    
    def _process_binary_file(self, file_path: Path) -> Document:
        """Process a binary file by extracting only metadata (no content).
        
        Args:
            file_path: Path to the binary file
            
        Returns:
            Processed Document with metadata but no content
        """
        doc_id = str(uuid.uuid4())
        
        # Extract basic file metadata
        extension = file_path.suffix.lower()
        doc_type = self._doc_type_mapping.get(extension, DocumentType.OTHER)
        
        # Create basic metadata
        metadata = UniversalMetadata(
            doc_id=doc_id,
            doc_type=doc_type,
            title=file_path.stem,
            file_path=str(file_path),
            storage_path=None,
            content_length=0,
            filename=file_path.name,
            extension=extension,
            size_bytes=file_path.stat().st_size,
            last_modified=datetime.fromtimestamp(file_path.stat().st_mtime),
            source_path=str(file_path),
            mime_type=mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
        )
        
        # Create document with empty content
        document = Document(
            doc_id=doc_id,
            content="",
            metadata=metadata,
            file_path=file_path
        )
        
        return document
    
    def _extract_basic_metadata(self, file_path: Path, content: str, doc_type: DocumentType) -> UniversalMetadata:
        """Extract basic metadata from a file.
        
        Args:
            file_path: Path to the file
            content: Extracted content
            doc_type: Document type
            
        Returns:
            Basic metadata
        """
        doc_id = str(uuid.uuid4())
        
        # Create basic metadata
        metadata = UniversalMetadata(
            doc_id=doc_id,
            doc_type=doc_type,
            title=file_path.stem,
            file_path=str(file_path),
            storage_path=None,
            content_length=len(content),
            filename=file_path.name,
            extension=file_path.suffix.lower(),
            size_bytes=file_path.stat().st_size,
            last_modified=datetime.fromtimestamp(file_path.stat().st_mtime),
            source_path=str(file_path),
            language='en',  # Default assumption
            mime_type=mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
        )
        
        return metadata
    
    def _extract_title_from_markdown(self, content: str) -> Optional[str]:
        """Extract title from markdown content.
        
        Args:
            content: Markdown content
            
        Returns:
            Extracted title if found, None otherwise
        """
        # Check for YAML frontmatter
        frontmatter_match = re.match(r'---\s+(.*?)\s+---', content, re.DOTALL)
        if frontmatter_match:
            frontmatter = frontmatter_match.group(1)
            title_match = re.search(r'^title:\s*(.*)$', frontmatter, re.MULTILINE)
            if title_match:
                return title_match.group(1).strip().strip('"\'')
        
        # Check for first heading
        heading_match = re.search(r'^#\s+(.*)$', content, re.MULTILINE)
        if heading_match:
            return heading_match.group(1).strip()
        
        return None
    
    def _extract_sections_from_markdown(self, content: str) -> List[Section]:
        """Extract sections from markdown content.
        
        Args:
            content: Markdown content
            
        Returns:
            List of sections
        """
        sections = []
        
        # Find all headings
        lines = content.splitlines()
        current_section = None
        current_level = 0
        current_content = []
        
        for line in lines:
            # Check if line is a heading
            heading_match = re.match(r'^(#+)\s+(.*)$', line)
            if heading_match:
                # If we have a current section, save it
                if current_section is not None:
                    sections.append(Section(
                        heading=current_section,
                        content="\n".join(current_content).strip(),
                        level=current_level
                    ))
                
                # Start new section
                current_level = len(heading_match.group(1))
                current_section = heading_match.group(2).strip()
                current_content = []
            else:
                # Add line to current section content
                if current_section is not None:
                    current_content.append(line)
        
        # Add the last section
        if current_section is not None:
            sections.append(Section(
                heading=current_section,
                content="\n".join(current_content).strip(),
                level=current_level
            ))
        
        return sections
    
    def _extract_sections_from_text(self, content: str) -> List[Section]:
        """Extract sections from plain text content.
        
        Args:
            content: Plain text content
            
        Returns:
            List of sections
        """
        sections = []
        
        # Try to identify sections based on common patterns
        # For example, all-caps lines followed by content
        lines = content.splitlines()
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Check if line might be a section heading (all caps, not too long)
            if line and line.isupper() and len(line) < 100 and len(line) > 3:
                # Found potential section heading
                heading = line
                section_content = []
                
                i += 1
                # Collect content until next potential heading or end
                while i < len(lines):
                    next_line = lines[i].strip()
                    if next_line and next_line.isupper() and len(next_line) < 100 and len(next_line) > 3:
                        # Found next heading
                        break
                    section_content.append(lines[i])
                    i += 1
                
                # Add section
                if section_content:
                    sections.append(Section(
                        heading=heading,
                        content="\n".join(section_content).strip(),
                        level=1
                    ))
            else:
                i += 1
        
        return sections
    
    def process_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> Document:
        """Process text and convert it to a document.
        
        Args:
            text: Text to process
            metadata: Optional metadata for the document
            
        Returns:
            Processed document
        """
        if not self._initialized:
            self.initialize({})
        
        # Generate a document ID
        doc_id = metadata.get('doc_id', str(uuid.uuid4())) if metadata else str(uuid.uuid4())
        
        # Create basic metadata
        universal_metadata = UniversalMetadata(
            doc_id=doc_id,
            doc_type=DocumentType.TEXT,
            title=metadata.get('title', 'Untitled Document') if metadata else 'Untitled Document',
            content_length=len(text),
            language=metadata.get('language', 'en') if metadata else 'en'
        )
        
        # Update with provided metadata
        if metadata:
            for key, value in metadata.items():
                if hasattr(universal_metadata, key):
                    try:
                        setattr(universal_metadata, key, value)
                    except Exception as e:
                        self._logger.warning(f"Could not set attribute {key}: {e}")
        
        # Try to extract sections
        if self._config.get('extract_sections', True):
            sections = self._extract_sections_from_text(text)
            if sections:
                universal_metadata.sections = sections
        
        # Create document
        document = Document(
            doc_id=doc_id,
            content=text,
            metadata=universal_metadata
        )
        
        return document
    
    def process_directory(self, directory_path: Path, 
                         recursive: bool = True, 
                         file_extensions: Optional[List[str]] = None) -> List[Document]:
        """Process all files in a directory and convert them to documents.
        
        Args:
            directory_path: Path to the directory to process
            recursive: Whether to process subdirectories
            file_extensions: Optional list of file extensions to process
            
        Returns:
            List of processed documents
        """
        if not self._initialized:
            self.initialize({})
        
        documents = []
        
        # Use the extensions we support by default if none specified
        if not file_extensions:
            file_extensions = list(self._supported_extensions.keys())
        
        # Normalize extensions
        extensions = [ext.lower() if ext.startswith('.') else f'.{ext.lower()}' for ext in file_extensions]
        
        # Get all matching files
        if recursive:
            for ext in extensions:
                for file_path in directory_path.glob(f'**/*{ext}'):
                    try:
                        document = self.process_file(file_path)
                        documents.append(document)
                    except Exception as e:
                        self._logger.error(f"Error processing file {file_path}: {e}")
        else:
            for ext in extensions:
                for file_path in directory_path.glob(f'*{ext}'):
                    try:
                        document = self.process_file(file_path)
                        documents.append(document)
                    except Exception as e:
                        self._logger.error(f"Error processing file {file_path}: {e}")
        
        return documents
