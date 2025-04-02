"""
Basic document processor implementation for the HADES ingest pipeline.

This module provides a basic implementation of the DocumentProcessor interface
that can handle common document formats like text, markdown, and PDF.
"""

import os
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import uuid

# Document format handlers
import markdown
import pypdf
import docx

from limnos.ingest.interface import DocumentProcessor, Document
from limnos.pipeline.interfaces import Configurable


class BasicDocumentProcessor(DocumentProcessor, Configurable):
    """Basic implementation of the DocumentProcessor interface."""
    
    def __init__(self):
        """Initialize the basic document processor."""
        self._logger = logging.getLogger(__name__)
        self._config = {}
        self._supported_extensions = {
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
        }
    
    def initialize(self, config: Dict[str, Any]) -> None:
        """Initialize the document processor with configuration.
        
        Args:
            config: Configuration dictionary
        """
        self._config = config
        
        # Add custom extensions if provided
        custom_extensions = config.get('custom_extensions', {})
        for ext, handler_name in custom_extensions.items():
            if hasattr(self, handler_name):
                self._supported_extensions[ext] = getattr(self, handler_name)
    
    @property
    def component_type(self) -> str:
        """Return the type of this component."""
        return "document_processor"
    
    @classmethod
    def get_plugin_type(cls) -> str:
        """Return the plugin type for this component."""
        return "document_processor"
    
    @classmethod
    def get_plugin_name(cls) -> str:
        """Return the name of this plugin."""
        return "basic"
    
    def process_file(self, file_path: Path) -> Document:
        """Process a file into a Document.
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Processed Document
            
        Raises:
            ValueError: If the file format is not supported
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        extension = file_path.suffix.lower()
        
        # Check if extension is supported
        if extension not in self._supported_extensions:
            raise ValueError(f"Unsupported file format: {extension}")
        
        # Process file based on extension
        handler = self._supported_extensions[extension]
        content = handler(file_path)
        
        # Create document
        doc_id = str(uuid.uuid4())
        metadata = {
            'source_path': str(file_path),
            'filename': file_path.name,
            'extension': extension,
            'size_bytes': file_path.stat().st_size,
            'last_modified': file_path.stat().st_mtime,
        }
        
        return Document(doc_id=doc_id, content=content, metadata=metadata)
    
    def process_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> Document:
        """Process raw text into a Document.
        
        Args:
            text: Raw text to process
            metadata: Optional metadata for the document
            
        Returns:
            Processed Document
        """
        doc_id = str(uuid.uuid4())
        metadata = metadata or {
            'source': 'raw_text',
        }
        
        return Document(doc_id=doc_id, content=text, metadata=metadata)
    
    def process_directory(self, directory_path: Path, recursive: bool = True) -> List[Document]:
        """Process all supported files in a directory.
        
        Args:
            directory_path: Path to the directory to process
            recursive: Whether to process subdirectories recursively
            
        Returns:
            List of processed Documents
        """
        directory_path = Path(directory_path)
        
        if not directory_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory_path}")
        
        documents = []
        
        # Get all files in directory
        if recursive:
            files = list(directory_path.glob('**/*'))
        else:
            files = list(directory_path.glob('*'))
        
        # Filter for regular files
        files = [f for f in files if f.is_file()]
        
        # Process each supported file
        for file_path in files:
            extension = file_path.suffix.lower()
            if extension in self._supported_extensions:
                try:
                    document = self.process_file(file_path)
                    documents.append(document)
                except Exception as e:
                    self._logger.error(f"Error processing file {file_path}: {e}")
        
        return documents
    
    def _process_text(self, file_path: Path) -> str:
        """Process a text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Extracted text content
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _process_markdown(self, file_path: Path) -> str:
        """Process a markdown file.
        
        Args:
            file_path: Path to the markdown file
            
        Returns:
            Extracted text content
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to plain text (simple approach)
        html_content = markdown.markdown(md_content)
        
        # Very basic HTML to text conversion
        # For production use, consider using a proper HTML-to-text converter
        text_content = html_content.replace('<p>', '').replace('</p>', '\n\n')
        text_content = text_content.replace('<h1>', '# ').replace('</h1>', '\n\n')
        text_content = text_content.replace('<h2>', '## ').replace('</h2>', '\n\n')
        text_content = text_content.replace('<h3>', '### ').replace('</h3>', '\n\n')
        text_content = text_content.replace('<h4>', '#### ').replace('</h4>', '\n\n')
        text_content = text_content.replace('<h5>', '##### ').replace('</h5>', '\n\n')
        text_content = text_content.replace('<h6>', '###### ').replace('</h6>', '\n\n')
        text_content = text_content.replace('<ul>', '').replace('</ul>', '\n')
        text_content = text_content.replace('<li>', '* ').replace('</li>', '\n')
        text_content = text_content.replace('<ol>', '').replace('</ol>', '\n')
        text_content = text_content.replace('<code>', '`').replace('</code>', '`')
        text_content = text_content.replace('<pre>', '```\n').replace('</pre>', '\n```\n')
        text_content = text_content.replace('<strong>', '**').replace('</strong>', '**')
        text_content = text_content.replace('<em>', '*').replace('</em>', '*')
        
        return text_content
    
    def _process_pdf(self, file_path: Path) -> str:
        """Process a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        text_content = ""
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n\n"
        
        except Exception as e:
            self._logger.error(f"Error extracting text from PDF {file_path}: {e}")
            text_content = f"[Error extracting PDF content: {str(e)}]"
        
        return text_content
    
    def _process_docx(self, file_path: Path) -> str:
        """Process a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        text_content = ""
        
        try:
            doc = docx.Document(file_path)
            
            for para in doc.paragraphs:
                text_content += para.text + "\n"
            
            # Add tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content += " | ".join(row_text) + "\n"
                text_content += "\n"
        
        except Exception as e:
            self._logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            text_content = f"[Error extracting DOCX content: {str(e)}]"
        
        return text_content
